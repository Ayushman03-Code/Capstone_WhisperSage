import openai
from docx import Document


OPENAI_API_KEY = ""


client = openai.OpenAI(api_key=OPENAI_API_KEY)

def generate_report(topic):
    """Generates a professional report using GPT-4o."""
    prompt = f"""
    Generate a detailed and professional report on the topic: {topic}. 
    The report should be well-structured and formatted using the following sections:

    **1. Title**
    - Example: *The Impact of Artificial Intelligence in Healthcare*

    **2. Introduction**
    - Provide an overview of the topic.
    - Example: *Artificial Intelligence (AI) is transforming the healthcare industry by improving diagnostics, patient care, and operational efficiency. This report explores the role of AI, its benefits, challenges, and future trends.*

    **3. Key Data and Statistics**
    - Present relevant facts and figures with sources if available.
    - Example:
        - *According to [source], AI in healthcare is projected to reach $45 billion by 2028.*
        - *A study by [organization] found that AI-driven diagnostics can improve accuracy by 30% compared to traditional methods.*

    **4. Benefits and Applications**
    - Discuss how the topic contributes to different areas.
    - Example: *AI helps in early disease detection, personalized medicine, and predictive analytics, leading to better patient outcomes.*

    **5. Challenges and Limitations**
    - Address potential concerns and obstacles.
    - Example: *Despite its benefits, AI faces challenges such as data privacy, regulatory issues, and the need for extensive training data.*

    **6. Real-World Examples / Case Studies**
    - Provide examples of how the topic is applied in real scenarios.
    - Example: *IBM Watson has been used to assist doctors in diagnosing cancer, reducing diagnosis time and improving accuracy.*

    **7. Future Trends and Recommendations**
    - Discuss emerging trends and offer suggestions.
    - Example: *AI is expected to revolutionize personalized medicine and robotic surgeries. Governments and institutions should invest in AI ethics and data security to maximize benefits.*

    **8. Conclusion**
    - Summarize the key points and final thoughts.
    - Example: *AI is reshaping healthcare, offering efficiency and accuracy. However, ethical and regulatory challenges must be addressed for widespread adoption.*

    Ensure the tone is formal, structured, and informative. Use clear and concise language, and maintain logical flow throughout the report.
    """

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}]
    )
    
    return response.choices[0].message.content  # Correct response parsing

def save_report_to_docx(report_text, filename="generated_report.docx"):
    """Saves the generated report as a .docx file."""
    doc = Document()
    doc.add_heading("Generated Report", level=1)
    
    for paragraph in report_text.split("\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph)
    
    doc.save(filename)
    print(f"✅ Report saved as {filename}")

if __name__ == "__main__":
    user_topic = """
    Market Growth & Adoption
    - The AI in healthcare market is projected to reach $187 billion by 2030, growing at a CAGR of 37% (Statista, 2023).
    - In 2023, over 60% of healthcare executives stated that AI will be critical to their operations by 2025 (PwC, 2023).
    - AI-powered chatbots and virtual assistants are projected to save the healthcare industry over $20 billion annually (Accenture, 2023).

    AI in Diagnostics & Treatment
    - AI-driven diagnostics improve accuracy by 30%, reducing misdiagnosis-related errors (Harvard Medical School, 2022).
    - AI algorithms can analyze medical images 1000x faster than humans, assisting in early disease detection (MIT, 2023).
    - Google’s DeepMind AI successfully detected over 50 eye diseases with an accuracy rate of 94% (Nature, 2022).

    Robotic Surgeries & AI-Assisted Procedures
    - The use of AI-assisted robotic surgery has increased by 32% in the last five years (JAMA, 2023).
    - AI-assisted robotic surgeries reduce post-operative complications by 21% (Mayo Clinic, 2023).

    Drug Discovery & Development
    - AI reduces drug discovery time by up to 70%, cutting costs by millions of dollars (Deloitte, 2023).
    - In 2021, the first AI-discovered drug entered human trials, speeding up the process by 5 years (Insilico Medicine, 2021).

    Administrative Efficiency & Cost Savings
    - AI can automate up to 80% of repetitive administrative tasks in hospitals, reducing costs by $18 billion annually (McKinsey, 2023).
    - Automated medical coding using AI can improve billing efficiency by 50% (Health Affairs, 2022).

    Patient Outcomes & Personalized Medicine
    - AI-driven personalized treatment plans can improve patient outcomes by 40% (Stanford Medicine, 2023).
    """

    report = generate_report(user_topic)
    save_report_to_docx(report)
