from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
import openai
from docx import Document
import os

app = FastAPI()


app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],  
    allow_credentials=True,
    allow_methods=["*"],  
    allow_headers=["*"],  
)


OPENAI_API_KEY = ""


client = openai.OpenAI(api_key=OPENAI_API_KEY)

def generate_report_content(user_input: str):
    """Generates a structured professional report using the original prompt."""
    prompt = f"""
    Generate a detailed and professional report on the topic: {user_input}. 
    The report should be well-structured and formatted using the following sections:

    **1. Title**
    - Example: *The Impact of Artificial Intelligence in Healthcare*

    **2. Introduction**
    - Provide an overview of the topic.
    - Example: *Artificial Intelligence (AI) is transforming the healthcare industry by improving diagnostics, patient care, and operational efficiency. This report explores the role of AI, its benefits, challenges, and future trends.*

    **3. Key Data and Statistics**
    - Present relevant facts and figures with sources if available.
    - Example:
        - *According to [source], AI in healthcare is projected to reach $45 billion by 2028.*
        - *A study by [organization] found that AI-driven diagnostics can improve accuracy by 30% compared to traditional methods.*

    **4. Benefits and Applications**
    - Discuss how the topic contributes to different areas.
    - Example: *AI helps in early disease detection, personalized medicine, and predictive analytics, leading to better patient outcomes.*

    **5. Challenges and Limitations**
    - Address potential concerns and obstacles.
    - Example: *Despite its benefits, AI faces challenges such as data privacy, regulatory issues, and the need for extensive training data.*

    **6. Real-World Examples / Case Studies**
    - Provide examples of how the topic is applied in real scenarios.
    - Example: *IBM Watson has been used to assist doctors in diagnosing cancer, reducing diagnosis time and improving accuracy.*

    **7. Future Trends and Recommendations**
    - Discuss emerging trends and offer suggestions.
    - Example: *AI is expected to revolutionize personalized medicine and robotic surgeries. Governments and institutions should invest in AI ethics and data security to maximize benefits.*

    **8. Conclusion**
    - Summarize the key points and final thoughts.
    - Example: *AI is reshaping healthcare, offering efficiency and accuracy. However, ethical and regulatory challenges must be addressed for widespread adoption.*

    Ensure the tone is formal, structured, and informative. Use clear and concise language, and maintain logical flow throughout the report.
    """

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}]
    )
    
    return response.choices[0].message.content

def save_report_to_docx(report_text: str, filename="generated_report.docx"):
    """Saves the report as a .docx file."""
    doc = Document()
    doc.add_heading("Generated Report", level=1)

    for paragraph in report_text.split("\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph)

    doc.save(filename)
    return filename

@app.post("/generate-report/")
async def generate_report(data: dict):
    """API Endpoint to generate and return a downloadable report."""
    topic = data.get("topic")
    if not topic:
        raise HTTPException(status_code=400, detail="Topic is required")

    try:
        report_text = generate_report_content(topic)
        filename = save_report_to_docx(report_text)

        return FileResponse(
            filename, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
            filename="generated_report.docx"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

